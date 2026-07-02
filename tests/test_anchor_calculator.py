from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document

from anchor_calculator import (
    AnchorInputs,
    CalculationError,
    calc_shear_breakout_geometry,
    calc_tension_breakout_geometry,
    calculate_anchor,
)
from report_writer import write_calculation_report


class AnchorCalculatorTests(unittest.TestCase):
    def test_default_values_match_manual_force_distribution_benchmark(self) -> None:
        result = calculate_anchor(AnchorInputs())
        v = result.values
        self.assertAlmostEqual(v["ca1"], 70.0, places=9)
        self.assertAlmostEqual(v["ca2"], 110.0, places=9)
        self.assertEqual(v["force_case"], "c")
        self.assertAlmostEqual(v["eccentricity"], 1250.0, places=9)
        self.assertAlmostEqual(v["l1"], 55.0, places=9)
        self.assertAlmostEqual(v["ae_star"], 606.798, places=9)
        self.assertAlmostEqual(v["modular_ratio"], 6.866666666666666, places=9)
        self.assertAlmostEqual(v["case_a_limit"], 61.666666666666664, places=9)
        self.assertAlmostEqual(v["case_b_limit"], 80.0, places=9)
        self.assertAlmostEqual(v["sigma_max"], 7.747653387888697, places=9)
        self.assertAlmostEqual(v["sigma_min"], 0.0, places=9)
        self.assertAlmostEqual(v["compression_x"], 92.90482726125902, places=9)
        self.assertAlmostEqual(v["total_anchor_tension"], 77.17224395705752, places=9)
        self.assertAlmostEqual(v["friction_force"], 38.86889758282301, places=9)
        self.assertAlmostEqual(v["nua"], 38.58612197852876, places=9)
        self.assertAlmostEqual(v["nuag"], 77.17224395705752, places=9)
        self.assertAlmostEqual(v["vua"], 0.0, places=9)
        self.assertAlmostEqual(v["vuag"], 0.0, places=9)
        self.assertAlmostEqual(v["nsa"], 107.70664500000001, places=9)
        self.assertAlmostEqual(v["h_eff_prime"], 220.0, places=9)
        self.assertAlmostEqual(v["tension_projected_width"], 400.0, places=9)
        self.assertAlmostEqual(v["tension_projected_length"], 400.0, places=9)
        self.assertAlmostEqual(v["anc"], 160000.0, places=9)
        self.assertAlmostEqual(v["anco"], 435600.0, places=9)
        self.assertAlmostEqual(v["ncbg"], 56.04906440643289, places=9)
        self.assertAlmostEqual(v["npn"], 63.20159999999999, places=9)
        self.assertAlmostEqual(v["vsa"], 51.6991896, places=9)
        self.assertAlmostEqual(v["shear_projected_width"], 390.0, places=9)
        self.assertAlmostEqual(v["shear_projected_depth"], 105.0, places=9)
        self.assertAlmostEqual(v["avc"], 40950.0, places=9)
        self.assertAlmostEqual(v["vcbg"], 23.658159928447517, places=9)
        self.assertAlmostEqual(v["vcpg"], 112.09812881286578, places=9)
        self.assertAlmostEqual(v["psi_ec_n"], 1.0, places=9)
        self.assertAlmostEqual(v["psi_ec_v"], 1.0, places=9)
        self.assertAlmostEqual(v["psi_h_v"], 1.0, places=9)
        self.assertAlmostEqual(v["kcp"], 2.0, places=9)
        self.assertAlmostEqual(v["grout_factor"], 0.8, places=9)
        self.assertAlmostEqual(result.tension_rebar_area, 257.2408131901917, places=9)
        self.assertAlmostEqual(v["tension_rebar_provided_area"], result.tension_rebar_area * 1.10, places=9)
        self.assertAlmostEqual(result.shear_rebar_area, 0.0, places=9)
        self.assertAlmostEqual(v["shear_rebar_provided_area"], 0.0, places=9)

        by_name = {check.name: check for check in result.checks}
        self.assertAlmostEqual(by_name["Steel strength in tension"].ratio, 0.47766934563822255, places=9)
        self.assertAlmostEqual(by_name["Concrete breakout strength in tension"].ratio, 1.8358259208385113, places=9)
        self.assertEqual(by_name["Concrete breakout strength in tension"].status, "Rebar Required")
        self.assertAlmostEqual(by_name["Pullout strength in tension"].ratio, 0.8140325978778757, places=9)
        self.assertIn("Minimum member thickness ha", by_name)
        self.assertEqual(by_name["Anchor shear demand after friction"].status, "Not Required")
        self.assertNotIn("Steel strength in shear", by_name)
        self.assertNotIn("Concrete breakout strength in shear", by_name)
        self.assertNotIn("Concrete pryout strength in shear", by_name)
        self.assertEqual(result.interaction_5_3.status, "Not Applicable")

    def test_tension_breakout_geometry_uses_projected_edges_and_effective_embedment_limit(self) -> None:
        default_geometry = calc_tension_breakout_geometry(AnchorInputs())
        self.assertEqual(default_geometry.close_edge_count, 4)
        self.assertTrue(default_geometry.h_eff_limited)
        self.assertAlmostEqual(default_geometry.h_eff_used, 220.0, places=9)
        self.assertAlmostEqual(default_geometry.projected_width, 400.0, places=9)
        self.assertAlmostEqual(default_geometry.projected_length, 400.0, places=9)
        self.assertAlmostEqual(default_geometry.area, 160000.0, places=9)

        large_pedestal = AnchorInputs(pedestal_l=1500.0, pedestal_b=900.0)
        large_geometry = calc_tension_breakout_geometry(large_pedestal)
        self.assertEqual(large_geometry.close_edge_count, 2)
        self.assertFalse(large_geometry.h_eff_limited)
        self.assertAlmostEqual(large_geometry.h_eff_used, large_pedestal.hef, places=9)
        self.assertAlmostEqual(large_geometry.area, 729000.0, places=9)

    def test_shear_breakout_geometry_limits_projected_width_for_wide_pedestals(self) -> None:
        geometry = calc_shear_breakout_geometry(AnchorInputs(pedestal_b=900.0))
        self.assertAlmostEqual(geometry.ca1_used, 70.0, places=9)
        self.assertAlmostEqual(geometry.projected_width, 390.0, places=9)
        self.assertAlmostEqual(geometry.projected_depth, 105.0, places=9)
        self.assertAlmostEqual(geometry.area, 40950.0, places=9)

        far_row = calc_shear_breakout_geometry(AnchorInputs(shear_case=2))
        self.assertAlmostEqual(far_row.ca1_used, 330.0, places=9)
        self.assertAlmostEqual(far_row.projected_width, 400.0, places=9)
        self.assertAlmostEqual(far_row.area, 198000.0, places=9)

    def test_minimum_spacing_equal_to_four_diameters_is_accepted(self) -> None:
        result = calculate_anchor(
            AnchorInputs(s1=88.0, s2=88.0, plate_l=198.0, plate_b=198.0, pedestal_l=300.0, pedestal_b=300.0, m=2.0)
        )
        by_name = {check.name: check for check in result.checks}
        self.assertEqual(by_name["Minimum spacing s1"].status, "OK")
        self.assertEqual(by_name["Minimum spacing s2"].status, "OK")
        self.assertIn(">=", by_name["Minimum spacing s1"].formula)
        self.assertIn(">=", by_name["Minimum spacing s2"].formula)

    def test_shear_demand_can_control_when_friction_is_exceeded(self) -> None:
        result = calculate_anchor(AnchorInputs(v=120.0, shear_case=2))
        self.assertGreater(result.values["vua"], 0.0)
        self.assertTrue(result.values["interaction_same_anchor_group"])
        self.assertTrue(result.values["interaction_applicable"])
        by_name = {check.name: check for check in result.checks}
        self.assertEqual(by_name["Anchor shear demand after friction"].status, "Required")
        self.assertIn("Steel strength in shear", by_name)
        self.assertIn("Concrete breakout strength in shear", by_name)
        self.assertIn("Concrete pryout strength in shear", by_name)
        self.assertEqual(result.interaction_5_3.section, "ACI 318-19 17.8.3")
        self.assertAlmostEqual(result.interaction_5_3.ratio, result.values["interaction_sum"] / 1.2, places=9)
        self.assertEqual(result.interaction_5_3.status, "NG")
        self.assertIn("> 1.2", result.interaction_5_3.substitution)

    def test_case1_shear_and_tension_on_different_rows_skips_interaction(self) -> None:
        result = calculate_anchor(AnchorInputs(v=120.0, shear_case=1))
        self.assertGreater(result.values["vua"], 0.0)
        self.assertEqual(result.values["tension_anchor_row"], "top")
        self.assertEqual(result.values["shear_resisting_row"], "bottom")
        self.assertFalse(result.values["interaction_same_anchor_group"])
        self.assertFalse(result.values["interaction_applicable"])
        self.assertIsNone(result.values["interaction_sum"])
        self.assertEqual(result.interaction_5_3.section, "ACI 318-19 17.8.1")
        self.assertEqual(result.interaction_5_3.status, "Not Applicable")
        self.assertIn("different anchor rows", result.interaction_5_3.note)

    def test_negative_moment_reverses_tension_row_and_case1_interaction(self) -> None:
        positive_case2 = calculate_anchor(AnchorInputs(m=25.0, v=120.0, shear_case=2))
        negative_case1 = calculate_anchor(AnchorInputs(m=-25.0, v=120.0, shear_case=1))

        self.assertAlmostEqual(negative_case1.values["eccentricity"], -positive_case2.values["eccentricity"], places=9)
        self.assertAlmostEqual(negative_case1.values["eccentricity_abs"], positive_case2.values["eccentricity_abs"], places=9)
        self.assertAlmostEqual(negative_case1.values["total_anchor_tension"], positive_case2.values["total_anchor_tension"], places=9)
        self.assertEqual(negative_case1.values["tension_anchor_row"], "bottom")
        self.assertEqual(negative_case1.values["shear_resisting_row"], "bottom")
        self.assertTrue(negative_case1.values["interaction_same_anchor_group"])
        self.assertTrue(negative_case1.values["interaction_applicable"])
        self.assertEqual(negative_case1.interaction_5_3.section, "ACI 318-19 17.8.3")
        self.assertAlmostEqual(negative_case1.values["interaction_sum"], positive_case2.values["interaction_sum"], places=9)

        negative_case2 = calculate_anchor(AnchorInputs(m=-25.0, v=120.0, shear_case=2))
        self.assertEqual(negative_case2.values["tension_anchor_row"], "bottom")
        self.assertEqual(negative_case2.values["shear_resisting_row"], "top")
        self.assertFalse(negative_case2.values["interaction_same_anchor_group"])
        self.assertEqual(negative_case2.interaction_5_3.status, "Not Applicable")

    def test_rebar_amplification_reduces_only_rebar_substituted_interaction_terms(self) -> None:
        baseline = calculate_anchor(AnchorInputs(v=100.0, shear_case=1))
        amplified = calculate_anchor(AnchorInputs(v=100.0, shear_case=1, tension_rebar_factor=1.5, shear_rebar_factor=1.5))
        self.assertAlmostEqual(baseline.values["tension_breakout_interaction_ratio"], 1.0 / 1.10, places=9)
        self.assertAlmostEqual(baseline.values["shear_breakout_interaction_ratio"], 1.0 / 1.10, places=9)
        self.assertAlmostEqual(amplified.values["tension_breakout_interaction_ratio"], 1.0 / 1.5, places=9)
        self.assertAlmostEqual(amplified.values["shear_breakout_interaction_ratio"], 1.0 / 1.5, places=9)
        amplified_checks = {check.name: check for check in amplified.checks}
        self.assertAlmostEqual(amplified.governing_tension_ratio, amplified_checks["Pullout strength in tension"].ratio, places=9)
        self.assertLess(amplified.values["shear_breakout_interaction_ratio"], baseline.values["shear_breakout_interaction_ratio"])
        self.assertAlmostEqual(amplified.governing_shear_ratio, amplified_checks["Steel strength in shear"].ratio, places=9)

    def test_manual_force_distribution_selects_three_eccentricity_cases(self) -> None:
        case_a = calculate_anchor(AnchorInputs(m=1.0))
        self.assertEqual(case_a.values["force_case"], "a")
        self.assertAlmostEqual(case_a.values["compression_x"], 370.0, places=9)
        self.assertAlmostEqual(case_a.values["total_anchor_tension"], 0.0, places=9)

        case_b = calculate_anchor(AnchorInputs(m=1.5))
        self.assertEqual(case_b.values["force_case"], "b")
        self.assertAlmostEqual(case_b.values["compression_x"], 330.0, places=9)
        self.assertAlmostEqual(case_b.values["total_anchor_tension"], 0.0, places=9)

        case_c = calculate_anchor(AnchorInputs(m=2.0))
        self.assertEqual(case_c.values["force_case"], "c")
        self.assertAlmostEqual(case_c.values["compression_x"], 268.6691181060678, places=9)
        self.assertAlmostEqual(case_c.values["total_anchor_tension"], 0.404213927590889, places=9)

    def test_built_up_grout_pad_controls_shear_steel_strength_factor(self) -> None:
        with_grout = calculate_anchor(AnchorInputs())
        without_grout = calculate_anchor(AnchorInputs(built_up_grout_pad=0.0))
        self.assertAlmostEqual(with_grout.values["grout_factor"], 0.8, places=9)
        self.assertAlmostEqual(without_grout.values["grout_factor"], 1.0, places=9)
        self.assertAlmostEqual(without_grout.values["vsa"], with_grout.values["vsa"] / 0.8, places=9)

    def test_invalid_geometry_raises_clear_error(self) -> None:
        with self.assertRaises(CalculationError):
            calculate_anchor(AnchorInputs(pedestal_l=250.0, s1=260.0))

    def test_invalid_hook_dimension_rejects_pullout_formula(self) -> None:
        with self.assertRaises(CalculationError):
            calculate_anchor(AnchorInputs(eh=50.0))

    def test_aci_scope_limits_are_enforced(self) -> None:
        with self.assertRaisesRegex(CalculationError, "17\\.3\\.1"):
            calculate_anchor(AnchorInputs(fc_prime=70.0))

        with self.assertRaisesRegex(CalculationError, "17\\.6\\.1\\.2"):
            calculate_anchor(AnchorInputs(futa=900.0))

        with self.assertRaisesRegex(CalculationError, "17\\.3\\.2"):
            calculate_anchor(AnchorInputs(da=102.0, eh=330.0))

    def test_word_report_is_created_and_readable(self) -> None:
        result = calculate_anchor(AnchorInputs())
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "anchor_report.docx"
            write_calculation_report(result, path, project_name="Unit Test Project", prepared_by="QA")
            self.assertTrue(path.exists())
            self.assertGreater(path.stat().st_size, 50_000)
            doc = Document(path)
            text = "\n".join(p.text for p in doc.paragraphs)
            self.assertIn("Anchor Bolt Design Calculation Report", text)
            self.assertIn("ACI 318-19 Chapter 17", text)
            self.assertGreaterEqual(len(doc.tables), 10)


if __name__ == "__main__":
    unittest.main()
