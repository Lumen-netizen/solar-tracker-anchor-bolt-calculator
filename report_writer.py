from __future__ import annotations

import math
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

from anchor_calculator import (
    ACI_MAX_ANCHOR_DIAMETER_MM,
    ACI_MAX_CAST_IN_FC_PRIME_MPA,
    ACI_MAX_FUTA_MPA,
    AnchorInputs,
    AnchorResults,
    CheckResult,
    INPUT_SPECS,
    SHEAR_CASE_LABELS,
)


REPORT_TITLE = "Anchor Bolt Design Calculation Report"

REPORT_CHECK_NAMES = {
    "Concrete local bearing": "Concrete local bearing",
    "Anchor shear demand after friction": "Anchor shear demand after friction",
    "Tension-shear interaction": "Tension and shear interaction",
}

PARAMETER_LABELS_EN = {
    "plate_l": "Base plate length in moment direction",
    "plate_b": "Base plate width perpendicular to moment direction",
    "da": "Anchor bolt diameter",
    "ase": "Effective steel area of one anchor",
    "hef": "Effective embedment depth",
    "s1": "Anchor row spacing in moment direction",
    "s2": "Anchor column spacing perpendicular to moment direction",
    "pedestal_l": "Concrete pedestal length in moment direction",
    "pedestal_b": "Concrete pedestal width perpendicular to moment direction",
    "ha": "Concrete member thickness / pedestal height",
    "eh": "Hook dimension for L- or J-bolt pullout",
    "built_up_grout_pad": "Built-up grout pad",
    "futa": "Specified tensile strength of anchor steel",
    "lambda_a": "Lightweight concrete modification factor",
    "fc_prime": "Specified compressive strength of concrete",
    "fc_design": "Design value for concrete local bearing",
    "ec_modulus": "Concrete elastic modulus",
    "es_modulus": "Steel elastic modulus",
    "fy_tension_rebar": "Design strength of tension anchor reinforcement",
    "fy_shear_rebar": "Design strength of shear anchor reinforcement",
    "tension_rebar_factor": "Provided tension reinforcement amplification factor",
    "shear_rebar_factor": "Provided shear reinforcement amplification factor",
    "n": "Factored axial compression",
    "m": "Factored base moment",
    "v": "Factored horizontal shear",
    "mu": "Interface friction coefficient",
    "phi_tension_steel": "Strength reduction factor for steel tension",
    "psi_ec_n": "Eccentricity factor for tension breakout",
    "psi_c_n": "Cracking factor for tension breakout",
    "psi_cp_n": "Splitting factor for tension breakout",
    "phi_tension_concrete": "Strength reduction factor for tension concrete breakout",
    "psi_c_p": "Cracking factor for pullout",
    "phi_pullout": "Strength reduction factor for pullout",
    "phi_shear_steel": "Strength reduction factor for steel shear",
    "psi_ec_v": "Eccentricity factor for shear breakout",
    "psi_c_v": "Cracking factor for shear breakout",
    "psi_h_v": "Member thickness factor for shear breakout",
    "phi_shear_concrete": "Strength reduction factor for shear concrete breakout",
    "kcp": "Concrete pryout coefficient",
    "phi_pryout": "Strength reduction factor for concrete pryout",
}


def write_calculation_report(
    results: AnchorResults,
    output_path: str | Path,
    project_name: str = "Photovoltaic Tracker Support",
    prepared_by: str = "",
) -> Path:
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _configure_document(doc)
    _add_cover(doc, project_name, prepared_by)
    _add_design_basis(doc, results)
    _add_layout_and_load_convention(doc, results)
    _add_input_summary(doc, results.inputs)
    _add_demand_section(doc, results)
    _add_checks_section(doc, results)
    _add_interaction_section(doc, results)
    _add_schematic_section(doc, results)
    _add_conclusion(doc, results)
    _add_appendix(doc, results)
    doc.save(output)
    return output


def _configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in (
        ("Heading 1", 15, "1F4E79"),
        ("Heading 2", 12, "1F4E79"),
        ("Heading 3", 10.5, "374151"),
    ):
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Anchor bolt design calculation - generated by the standalone design app")
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(100, 116, 139)


def _add_cover(doc: Document, project_name: str, prepared_by: str) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(REPORT_TITLE)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(project_name or "Photovoltaic Tracker Support")
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = RGBColor(71, 85, 105)

    info = [
        ("Anchor design code", "ACI 318-19 Chapter 17"),
        ("Force distribution basis", "Table 8-3 of 《钢结构节点设计手册》（第四版）"),
        ("Connection type", "Four cast-in L-bolts anchoring a base plate to a concrete pedestal"),
        ("Report date", datetime.now().strftime("%Y-%m-%d")),
        ("Prepared by", prepared_by or "Engineer"),
    ]
    table = _add_table(doc, ["Item", "Description"], info, widths=(1.8, 5.6))
    _shade_header(table)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    note.add_run("Note. ").bold = True
    note.add_run(
        "This report is generated from a standalone calculation engine. The original spreadsheet was used only as a benchmark for the initial implementation."
    )


def _add_design_basis(doc: Document, results: AnchorResults) -> None:
    doc.add_heading("1. Design Basis and Scope", level=1)
    bullets = [
        "ACI 318-19 Chapter 17 is used for anchor strength checks in tension, shear, and combined tension and shear.",
        "Base plate compression stress, compression depth, and anchor tension demand are calculated using Table 8-3 of 《钢结构节点设计手册》（第四版）. The source title is retained in Chinese for traceability.",
        "The calculation covers a four-anchor rectangular layout for a photovoltaic tracker support base plate.",
        "The demand model assumes uniaxial base moment and horizontal shear along the selected shear case direction. Biaxial moment, torsion, and arbitrary anchor coordinates are outside the current scope.",
        "Units are mm, MPa, kN, and kN.m. Strength ratios are calculated as demand divided by design strength.",
        "The anchor type is treated as a cast-in L-bolt. The ACI pullout expression for J- or L-bolts is used for pullout strength.",
        f"ACI 318-19 scope limits are enforced by the program: f'c for cast-in anchor calculations <= {ACI_MAX_CAST_IN_FC_PRIME_MPA:.1f} MPa per 17.3.1; futa <= {ACI_MAX_FUTA_MPA:.1f} MPa per 17.6.1.2 and 17.7.1.2; da <= {ACI_MAX_ANCHOR_DIAMETER_MM:.1f} mm per 17.3.2. The user shall also confirm futa <= 1.9fya because fya is not a program input.",
        "If interface friction is not sufficient for the applied horizontal shear, the residual shear is assigned to the two shear-resisting anchors and checked by ACI 318-19 Chapter 17.",
        "Adhesive anchor bond failure is not applicable to this cast-in L-bolt model. Seismic-specific anchorage requirements, side-face blowout, and arbitrary anchor coordinates are outside the current scope.",
    ]
    _add_bullets(doc, bullets)
    if results.warnings:
        _add_note_box(doc, "Calculation Notes", results.warnings)


def _add_layout_and_load_convention(doc: Document, results: AnchorResults) -> None:
    doc.add_heading("2. Anchor Layout and Load Convention", level=1)
    diagram_path = _make_report_diagram(results)
    doc.add_picture(str(diagram_path), width=Inches(6.7))
    caption = doc.add_paragraph("Figure 1. Anchor layout, edge distances, embedment, and positive load convention.")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in caption.runs:
        run.font.size = Pt(8.5)
        run.font.italic = True
        run.font.color.rgb = RGBColor(71, 85, 105)

    rows = [
        ("Axial force N", "Positive N is taken as downward compression on the base plate."),
        ("Moment M", "Positive M produces tension in the anchor row shown in the plan diagram."),
        ("Shear V", "Positive V follows the blue arrow direction shown in the plan diagram."),
        ("Shear case", SHEAR_CASE_LABELS[results.inputs.shear_case]),
    ]
    table = _add_table(doc, ["Load item", "Convention"], rows, widths=(1.8, 5.6))
    _shade_header(table)


def _add_input_summary(doc: Document, inputs: AnchorInputs) -> None:
    doc.add_heading("3. Input Data", level=1)
    for group in ("Geometry", "Materials", "Loads", "Factors"):
        doc.add_heading(group, level=2)
        rows = []
        for spec in INPUT_SPECS:
            if spec.group != group:
                continue
            value = getattr(inputs, spec.key)
            if spec.key == "built_up_grout_pad":
                value_text = "Yes" if value >= 0.5 else "No"
            else:
                value_text = _fmt(value, spec.decimals)
            rows.append((_parameter_name(spec.key), spec.symbol, value_text, spec.unit))
        if group == "Loads":
            rows.append(("Shear control case", "Case", str(inputs.shear_case), "-"))
        table = _add_table(doc, ["Parameter", "Symbol", "Value", "Unit"], rows, widths=(2.6, 1.1, 1.5, 1.0))
        _shade_header(table)


def _add_demand_section(doc: Document, results: AnchorResults) -> None:
    doc.add_heading("4. Demand Calculation", level=1)
    v = results.values
    inputs = results.inputs
    p = doc.add_paragraph()
    p.add_run("Force distribution basis. ").bold = True
    p.add_run(
        "The base plate compression block and anchor tension demand are evaluated using Table 8-3 of "
        "《钢结构节点设计手册》（第四版）. Three eccentricity cases are considered by the program."
    )
    case_rows = [
        ("Case a", "e <= l/6", "Full base plate compression; Ta = 0."),
        ("Case b", "l/6 < e <= l/6 + l1/3", "Triangular compression block; Ta = 0."),
        ("Case c", "e > l/6 + l1/3", "Partial compression with anchor tension; xn is solved from Eq. 8-158."),
    ]
    table = _add_table(doc, ["Case", "Criterion", "Meaning"], case_rows, widths=(0.9, 2.2, 4.3))
    _shade_header(table)

    n_eq = (
        f"({inputs.es_modulus:g} x 10^5) / ({inputs.ec_modulus:g} x 10^4)"
        if inputs.ec_modulus > 0
        else "-"
    )
    rows = [
        ("Edge distance ca1", "ca1 = (L - s1) / 2", f"({inputs.pedestal_l:g} - {inputs.s1:g}) / 2", f"{v['ca1']:.3f} mm"),
        ("Edge distance ca2", "ca2 = (B - s2) / 2", f"({inputs.pedestal_b:g} - {inputs.s2:g}) / 2", f"{v['ca2']:.3f} mm"),
        ("Distance from tension edge to anchor row", "l1 = (l - s1) / 2", f"({inputs.plate_l:g} - {inputs.s1:g}) / 2", f"{v['l1']:.3f} mm"),
        ("Load eccentricity", "e = M / N", f"{inputs.m:g} / {inputs.n:g} x 1000", f"{v['eccentricity']:.3f} mm"),
        ("Elastic modular ratio", "n = Es / Ec", n_eq, f"{v['modular_ratio']:.3f}"),
        ("Total effective area of tension anchors", "A*e = 2 Ase", f"2 x {inputs.ase:g}", f"{v['ae_star']:.3f} mm2"),
        ("Case a limit", "l / 6", f"{inputs.plate_l:g} / 6", f"{v['case_a_limit']:.3f} mm"),
        ("Case b limit", "l / 6 + l1 / 3", f"{v['case_a_limit']:.3f} + {v['l1']:.3f} / 3", f"{v['case_b_limit']:.3f} mm"),
        ("Selected force case", "Table 8-3 criterion", f"e = {v['eccentricity']:.3f} mm", f"{v['force_case']} - {v['force_case_label']}"),
        ("Compression block depth", "xn by selected case", _force_depth_substitution(results), f"{v['compression_x']:.3f} mm"),
        ("Maximum concrete compression stress", v["bearing_formula"], v["bearing_substitution"], f"{v['sigma_max']:.3f} MPa"),
        ("Total tension force in tension-side anchors", "Ta by selected case", _force_tension_substitution(results), f"{v['total_anchor_tension']:.3f} kN"),
        ("Friction force", "Vfb = mu(Ta + N)", f"{inputs.mu:g} x ({v['total_anchor_tension']:.3f} + {inputs.n:g})", f"{v['friction_force']:.3f} kN"),
        ("Anchor group tension demand", "Nua,g = Ta", f"{v['total_anchor_tension']:.3f}", f"{v['nuag']:.3f} kN"),
        ("Single-anchor tension demand", "Nua = Nua,g / 2", f"{v['nuag']:.3f} / 2", f"{v['nua']:.3f} kN"),
        ("Anchor group shear demand after friction", "Vua,g = max(V - Vfb, 0)", f"max({inputs.v:g} - {v['friction_force']:.3f}, 0)", f"{v['vuag']:.3f} kN"),
        ("Single-anchor shear demand", "Vua = max((V - Vfb) / 2, 0)", f"max(({inputs.v:g} - {v['friction_force']:.3f}) / 2, 0)", f"{v['vua']:.3f} kN"),
    ]
    table = _add_table(doc, ["Item", "Formula", "Substitution", "Result"], rows, widths=(1.9, 2.35, 2.15, 1.25))
    _shade_header(table)


def _add_checks_section(doc: Document, results: AnchorResults) -> None:
    doc.add_heading("5. Strength Checks", level=1)
    checks_by_name = {check.name: check for check in results.checks}
    groups = [
        (
            "5.1 Geometry Requirements and Concrete Local Bearing",
            ["Minimum spacing s1", "Minimum spacing s2", "Minimum member thickness ha", "Concrete local bearing"],
        ),
        (
            "5.2 Tension Checks",
            ["Steel strength in tension", "Concrete breakout strength in tension", "Pullout strength in tension"],
        ),
        (
            "5.3 Shear Checks",
            [
                "Anchor shear demand after friction",
                "Steel strength in shear",
                "Concrete breakout strength in shear",
                "Concrete pryout strength in shear",
            ],
        ),
    ]
    for title, names in groups:
        checks = [checks_by_name[name] for name in names if name in checks_by_name]
        doc.add_heading(title, level=2)
        _add_check_table(doc, checks)
        if title.endswith("Tension Checks"):
            if results.tension_rebar_area > 0:
                doc.add_paragraph(
                    "The required tension anchor reinforcement area is calculated by assigning the full anchor group "
                    f"tension demand to reinforcement: As,N,req = {results.tension_rebar_area:.2f} mm2. "
                    f"For the interaction check, the provided area is taken as As,N,prov = As,N,req x k_As,N = "
                    f"{results.values['tension_rebar_provided_area']:.2f} mm2."
                )
            else:
                doc.add_paragraph("No tension anchor reinforcement is required by the concrete breakout check in this calculation.")
        if title.endswith("Shear Checks"):
            if results.values["vua"] <= 1.0e-9:
                doc.add_paragraph(
                    "Anchor shear strength checks are not required because the interface friction force is not less than the applied shear."
                )
            else:
                doc.add_paragraph(
                    "The residual shear after interface friction is assigned to the two shear-resisting anchors. "
                    f"When concrete shear breakout requires anchor reinforcement, the required shear reinforcement area is "
                    f"As,V,req = {results.shear_rebar_area:.2f} mm2 and the provided area used for interaction is "
                    f"As,V,prov = {results.values['shear_rebar_provided_area']:.2f} mm2."
                )


def _add_interaction_section(doc: Document, results: AnchorResults) -> None:
    doc.add_heading("6. Tension-Shear Interaction", level=1)
    v = results.values
    if v["vua"] <= 1.0e-9:
        current_result = "Not applicable because anchor shear demand after interface friction is zero."
    elif not v.get("interaction_same_anchor_group", True):
        current_result = (
            "Not applicable because the shear-resisting anchors and the tension anchors are different rows "
            "for the selected shear case."
        )
    elif results.governing_tension_ratio <= 0.2 or results.governing_shear_ratio <= 0.2:
        current_result = (
            "Interaction may be neglected because at least one ratio is not greater than 0.2 "
            f"(eta_N = {results.governing_tension_ratio:.3f}, eta_V = {results.governing_shear_ratio:.3f})."
        )
    else:
        interaction_sum = v.get("interaction_sum") or 0.0
        interaction_limit = v.get("interaction_limit", 1.2)
        status = "OK" if interaction_sum <= interaction_limit else "NG"
        current_result = f"eta_N + eta_V = {interaction_sum:.3f} <= {interaction_limit:.1f}; status = {status}."
    rows = [
        ("ACI 318-19 17.8.1", "Anchors or anchor groups resisting both tension and shear shall satisfy 17.8.2 and 17.8.3."),
        ("ACI 318-19 17.8.2(a)", "The interaction between tension and shear may be neglected if Nua/(phi Nn) <= 0.2."),
        ("ACI 318-19 17.8.2(b)", "The interaction between tension and shear may be neglected if Vua/(phi Vn) <= 0.2."),
        ("ACI 318-19 17.8.3", "If both ratios are greater than 0.2, check Nua/(phi Nn) + Vua/(phi Vn) <= 1.2."),
        ("Current eta_N", f"{results.governing_tension_ratio:.3f}"),
        ("Current eta_V", f"{results.governing_shear_ratio:.3f}"),
        ("Current result", current_result),
    ]
    table = _add_table(doc, ["Reference / item", "Requirement or result"], rows, widths=(2.0, 5.4))
    _shade_header(table)


def _add_schematic_section(doc: Document, results: AnchorResults) -> None:
    doc.add_heading("7. Failure Mode and Reinforcement Schematics", level=1)
    p = doc.add_paragraph()
    p.add_run("Schematic figures. ").bold = True
    p.add_run("The following simplified figures are generated for engineering interpretation of the check items and are not reproduced code figures.")
    sheet_path = _make_failure_schematic_sheet(results)
    doc.add_picture(str(sheet_path), width=Inches(6.7))


def _add_conclusion(doc: Document, results: AnchorResults) -> None:
    doc.add_heading("8. Governing Results and Engineering Conclusion", level=1)
    status_text = {
        "OK": "The anchor connection satisfies the checks covered by this calculation.",
        "Rebar Required": "Additional anchor reinforcement is required for at least one concrete failure mode covered by this calculation.",
        "NG": "At least one covered check does not satisfy the required strength criterion. Revise the anchor size, embedment, pedestal dimensions, reinforcement, or applied loads.",
    }[results.overall_status]
    rows = [
        ("Overall status", results.overall_status),
        ("Governing tension utilization", f"{results.governing_tension_ratio:.3f}"),
        ("Governing shear utilization", f"{results.governing_shear_ratio:.3f}"),
        ("Required tension reinforcement", f"As,N,req = {results.tension_rebar_area:.2f} mm2"),
        ("Provided tension reinforcement used for interaction", f"As,N,prov = {results.values['tension_rebar_provided_area']:.2f} mm2"),
        ("Required shear reinforcement", f"As,V,req = {results.shear_rebar_area:.2f} mm2"),
        ("Provided shear reinforcement used for interaction", f"As,V,prov = {results.values['shear_rebar_provided_area']:.2f} mm2"),
        ("Conclusion", status_text),
    ]
    table = _add_table(doc, ["Item", "Result"], rows, widths=(2.3, 5.0))
    _shade_header(table)
    _shade_status_cells(table)


def _add_appendix(doc: Document, results: AnchorResults) -> None:
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("Appendix A. Formula Trace", level=1)
    rows = []
    for key, label, unit in (
        ("l1", "Distance from tension edge to anchor row l1", "mm"),
        ("ae_star", "Total effective area of tension anchors A*e", "mm2"),
        ("modular_ratio", "Elastic modular ratio n", "-"),
        ("case_a_limit", "Case a eccentricity limit", "mm"),
        ("case_b_limit", "Case b eccentricity limit", "mm"),
        ("compression_x", "Compression block depth xn", "mm"),
        ("sigma_max", "Maximum concrete compression stress", "MPa"),
        ("total_anchor_tension", "Total tension force in tension-side anchors Ta", "kN"),
        ("friction_force", "Interface friction force Vfb", "kN"),
        ("nuag", "Required tension strength of anchor group Nua,g", "kN"),
        ("nua", "Required tension strength of single anchor Nua", "kN"),
        ("vuag", "Required shear strength of anchor group Vua,g", "kN"),
        ("vua", "Required shear strength of single anchor Vua", "kN"),
        ("nsa", "Nominal steel tension strength Nsa", "kN"),
        ("h_eff_prime", "Limited effective embedment h'ef", "mm"),
        ("tension_close_edge_count", "Number of edges within 1.5h'ef for tension breakout", "-"),
        ("tension_ca_near", "Near edge distance for tension breakout", "mm"),
        ("tension_ca_far", "Far edge distance for tension breakout", "mm"),
        ("tension_projection", "Tension breakout projection distance 1.5h'ef", "mm"),
        ("tension_projected_width", "Projected tension breakout width", "mm"),
        ("tension_projected_length", "Projected tension breakout length", "mm"),
        ("anc", "Projected tension breakout area ANc", "mm2"),
        ("anco", "Reference projected tension breakout area ANco", "mm2"),
        ("tension_capped_area", "Capped projected tension breakout area used in strength equation", "mm2"),
        ("nb", "Basic tension breakout strength Nb", "kN"),
        ("psi_ed_n", "Tension edge effect factor psi_ed,N", "-"),
        ("ncbg", "Nominal tension breakout strength Ncbg", "kN"),
        ("npn", "Nominal pullout strength Npn", "kN"),
        ("vsa", "Nominal steel shear strength Vsa", "kN"),
        ("shear_ca1_actual", "Actual shear edge distance ca1", "mm"),
        ("shear_ca1", "Shear edge distance ca1 used in ACI 17.7.2", "mm"),
        ("critical_spacing", "Critical shear spacing", "mm"),
        ("shear_projection", "Shear breakout projection depth 1.5ca1", "mm"),
        ("shear_projected_width", "Projected shear breakout width", "mm"),
        ("shear_projected_depth", "Projected shear breakout depth", "mm"),
        ("avc", "Projected shear breakout area AVc", "mm2"),
        ("avco", "Reference projected shear breakout area AVco", "mm2"),
        ("shear_capped_area", "Capped projected shear breakout area used in strength equation", "mm2"),
        ("vb", "Basic shear breakout strength Vb", "kN"),
        ("psi_ed_v", "Shear edge effect factor psi_ed,V", "-"),
        ("vcbg", "Nominal shear breakout strength Vcbg", "kN"),
        ("vcpg", "Nominal pryout strength Vcpg", "kN"),
        ("tension_rebar_provided_area", "Provided tension anchor reinforcement area As,N,prov", "mm2"),
        ("shear_rebar_provided_area", "Provided shear anchor reinforcement area As,V,prov", "mm2"),
        ("tension_breakout_interaction_ratio", "Tension breakout ratio used for interaction", "-"),
        ("shear_breakout_interaction_ratio", "Shear breakout ratio used for interaction", "-"),
    ):
        rows.append((label, key, _fmt(results.values[key], 3), unit))
    table = _add_table(doc, ["Quantity", "Variable", "Value", "Unit"], rows, widths=(3.4, 1.4, 1.4, 0.8))
    _shade_header(table)

    doc.add_heading("Appendix B. Limitations", level=1)
    _add_bullets(
        doc,
        [
            "The calculation does not replace project-specific engineering judgment.",
            "The user should confirm that the selected anchor type, material properties, and load combinations are applicable to the project.",
            "Where reinforcement is required, anchorage and detailing of that reinforcement should be designed and checked separately.",
            "The report references ACI section numbers for traceability and does not reproduce the full code text.",
        ],
    )


def _add_check_table(doc: Document, checks: Iterable[CheckResult]) -> None:
    rows = []
    for check in checks:
        if check.ratio is None:
            ratio = "-"
        elif check.name == "Anchor shear demand after friction":
            ratio = f"V/Vfb = {check.ratio:.3f}"
        elif check.status in {"OK", "NG", "Rebar Required"}:
            operator = "<=" if check.ratio <= 1.0 else ">"
            ratio = f"{check.ratio:.3f} {operator} 1.0"
        else:
            ratio = f"{check.ratio:.3f}"
        demand = "-" if check.demand is None else f"{check.demand:.3f}"
        capacity = "-" if check.capacity is None else f"{check.capacity:.3f}"
        rows.append((_report_check_name(check.name), check.section, check.formula, check.substitution, demand, capacity, ratio, check.status))
    table = _add_table(
        doc,
        ["Check", "Reference", "Formula", "Substitution", "Demand", "Design strength", "Ratio", "Status"],
        rows,
        widths=(1.45, 1.0, 1.5, 2.0, 0.75, 0.9, 0.65, 0.75),
    )
    _shade_header(table)
    _shade_status_cells(table)


def _report_check_name(name: str) -> str:
    return REPORT_CHECK_NAMES.get(name, name)


def _parameter_name(key: str) -> str:
    return PARAMETER_LABELS_EN.get(key, key)


def _force_depth_substitution(results: AnchorResults) -> str:
    v = results.values
    i = results.inputs
    if v["force_case"] == "a":
        return f"x = l = {i.plate_l:g}"
    if v["force_case"] == "b":
        return f"x = 3(l/2 - e) = 3({i.plate_l:g}/2 - {v['eccentricity']:.3f})"
    return (
        "xn solved from Eq. 8-158: "
        "xn^3 + 3(e - l/2)xn^2 - (6nA*e/b)(e + l/2 - l1)(l - l1 - xn) = 0"
    )


def _force_tension_substitution(results: AnchorResults) -> str:
    v = results.values
    i = results.inputs
    if v["force_case"] in {"a", "b"}:
        return "Ta = 0 for the selected force case"
    return (
        f"{i.n:g} x ({v['eccentricity']:.3f} - {i.plate_l:g}/2 + {v['compression_x']:.3f}/3) / "
        f"({i.plate_l:g} - {v['l1']:.3f} - {v['compression_x']:.3f}/3)"
    )


def _add_table(doc: Document, headers: list[str], rows: Iterable[tuple], widths: tuple[float, ...]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        _set_cell_width(cell, widths[idx])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8.5)
    for row_data in rows:
        row = table.add_row()
        for idx, item in enumerate(row_data):
            cell = row.cells[idx]
            cell.text = str(item)
            _set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx >= len(row_data) - 3 else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    doc.add_paragraph()
    return table


def _set_cell_width(cell, width_in: float) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def _shade_header(table) -> None:
    for cell in table.rows[0].cells:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "1F4E79")
        cell._tc.get_or_add_tcPr().append(shading)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)


def _shade_status_cells(table) -> None:
    for row in table.rows[1:]:
        for cell in row.cells:
            text = cell.text.strip()
            fill = None
            if text == "OK":
                fill = "DFF4E8"
            elif text == "NG":
                fill = "FEE2E2"
            elif text == "Rebar Required":
                fill = "FEF3C7"
            elif text == "Not Applicable":
                fill = "E5E7EB"
            elif text in {"Required", "Not Required"}:
                fill = "DBEAFE"
            if fill:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), fill)
                cell._tc.get_or_add_tcPr().append(shading)


def _add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style=None)
        p.style = doc.styles["List Bullet"]
        p.add_run(item)


def _add_note_box(doc: Document, title: str, notes: Iterable[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "EFF6FF")
    cell._tc.get_or_add_tcPr().append(shading)
    p = cell.paragraphs[0]
    p.add_run(title + ": ").bold = True
    p.add_run(" ".join(notes))
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(8.8)
    doc.add_paragraph()


def _make_interaction_diagram(results: AnchorResults) -> Path:
    path = Path(tempfile.gettempdir()) / "anchor_bolt_interaction_diagram.png"
    img = Image.new("RGB", (760, 430), "white")
    draw = ImageDraw.Draw(img)
    font = _font(22)
    small = _font(18)
    _draw_interaction_plot(draw, 95, 52, 560, 280, results, font, small)
    img.save(path)
    return path


def _make_failure_schematic_sheet(results: AnchorResults) -> Path:
    path = Path(tempfile.gettempdir()) / "anchor_bolt_failure_schematics.png"
    img = Image.new("RGB", (1320, 1120), "white")
    draw = ImageDraw.Draw(img)
    title_font = _font(28)
    card_font = _font(20)
    small = _font(17)
    draw.text((42, 30), "Simplified Failure Mode and Reinforcement Schematics", fill="#111827", font=title_font)

    cards = [
        ("Concrete local bearing", "bearing"),
        ("Steel failure in tension", "steel_tension"),
        ("Pullout of J- or L-bolt", "pullout"),
        ("Concrete breakout in tension", "tension_breakout"),
        ("Steel failure in shear", "steel_shear"),
        ("Concrete breakout in shear", "shear_breakout"),
        ("Concrete pryout in shear", "pryout"),
        ("Interface friction check", "friction"),
    ]
    x0, y0 = 34, 84
    card_w, card_h = 610, 236
    gap_x, gap_y = 28, 22
    for idx, (title, mode) in enumerate(cards):
        col = idx % 2
        row = idx // 2
        x = x0 + col * (card_w + gap_x)
        y = y0 + row * (card_h + gap_y)
        _draw_schematic_card(draw, x, y, card_w, card_h, title, mode, results, card_font, small)
    img.save(path)
    return path


def _draw_schematic_card(
    draw: ImageDraw.ImageDraw,
    x: float,
    y: float,
    w: float,
    h: float,
    title: str,
    mode: str,
    results: AnchorResults,
    font,
    small,
) -> None:
    black = "#111827"
    draw.rounded_rectangle((x, y, x + w, y + h), radius=8, fill="#F8FAFC", outline="#CBD5E1", width=2)
    draw.text((x + 18, y + 14), title, fill=black, font=font)

    if mode == "interaction":
        _draw_interaction_plot(draw, x + 34, y + 58, w - 68, h - 84, results, font, small, show_point=True)
        return

    if mode == "steel_tension":
        _draw_report_anchor_section(draw, x + 70, y + 68, 220, 130, "steel_tension", results, small)
        draw.text((x + 360, y + 96), "Steel failure of anchor\nunder tensile demand.", fill="#475569", font=small)
    elif mode == "pullout":
        _draw_report_anchor_section(draw, x + 70, y + 68, 240, 130, "pullout", results, small)
        draw.text((x + 358, y + 92), "Void after pullout is shown\nwith the same width as the\nanchor shank.", fill="#475569", font=small)
    elif mode == "tension_breakout":
        _draw_report_tension_breakout(draw, x + 34, y + 62, 530, 142, results, small)
    elif mode == "steel_shear":
        _draw_report_anchor_section(draw, x + 70, y + 66, 230, 136, "steel_shear", results, small)
        draw.text((x + 360, y + 100), "Steel shear failure near\nthe concrete surface.", fill="#475569", font=small)
    elif mode == "shear_breakout":
        _draw_report_shear_breakout(draw, x + 34, y + 62, 540, 142, results, small)
    elif mode == "pryout":
        _draw_report_anchor_section(draw, x + 64, y + 64, 260, 140, "pryout", results, small)
        draw.text((x + 380, y + 96), "Large pryout cone starts\nnear the anchor bottom;\nsmall bearing cone is at\nthe loaded side.", fill="#475569", font=small)
    elif mode == "bearing":
        _draw_report_local_bearing(draw, x + 58, y + 64, 280, 136, results, small)
    elif mode == "friction":
        _draw_report_friction(draw, x + 58, y + 70, 310, 124, results, small)


def _draw_report_bolt_with_nut(
    draw: ImageDraw.ImageDraw,
    x: float,
    plate_top_y: float,
    rod_bottom_y: float,
    *,
    rod_color: str = "#F97316",
    outline: str = "#111827",
    rod_width: int = 8,
) -> None:
    draw.line((x, plate_top_y - 26, x, rod_bottom_y), fill=rod_color, width=rod_width)
    draw.rectangle((x - 5, plate_top_y - 34, x + 5, plate_top_y - 18), fill=rod_color, outline=outline)
    draw.rectangle((x - 11, plate_top_y - 18, x + 11, plate_top_y - 6), fill=rod_color, outline=outline)
    draw.polygon(
        [
            (x - 15, plate_top_y - 6),
            (x + 15, plate_top_y - 6),
            (x + 11, plate_top_y + 2),
            (x - 11, plate_top_y + 2),
        ],
        fill=rod_color,
        outline=outline,
    )


def _draw_report_single_base(draw: ImageDraw.ImageDraw, x: float, y: float, bw: float, bh: float) -> tuple[float, float, float, float]:
    black = "#111827"
    concrete = "#D1D5DB"
    plate = "#94A3B8"
    concrete_top = y + 38
    concrete_bottom = y + bh
    plate_top = y + 24
    plate_bottom = concrete_top
    draw.rectangle((x, concrete_top, x + bw, concrete_bottom), fill=concrete, outline=black, width=2)
    draw.rectangle((x + bw * 0.28, plate_top, x + bw * 0.72, plate_bottom), fill=plate, outline=black, width=1)
    return concrete_top, concrete_bottom, plate_top, plate_bottom


def _draw_report_anchor_section(
    draw: ImageDraw.ImageDraw,
    x: float,
    y: float,
    bw: float,
    bh: float,
    mode: str,
    results: AnchorResults,
    small,
) -> None:
    black = "#111827"
    blue = "#2563EB"
    green = "#16A34A"
    orange = "#F97316"
    grey = "#9CA3AF"
    red = "#DC2626"
    cyan = "#06B6D4"

    if mode == "pullout":
        concrete_top = y + 46
        concrete_bottom = y + bh
        ax = x + bw * 0.5
        plate_top = y + 18
        plate_bottom = y + 32
        draw.rectangle((x, concrete_top, x + bw, concrete_bottom), fill="#D1D5DB", outline=black, width=2)
        draw.rectangle((x + bw * 0.28, plate_top, x + bw * 0.72, plate_bottom), fill="#94A3B8", outline=black, width=1)
        void_half = 4
        void_top = y + 90
        void_bottom = min(concrete_bottom - 8, y + 122)
        draw.line((ax, y + 4, ax, void_top), fill=orange, width=8)
        draw.rectangle((ax - 5, y - 22, ax + 5, y - 6), fill=orange, outline=black)
        draw.rectangle((ax - 11, y - 6, ax + 11, y + 6), fill=orange, outline=black)
        draw.polygon([(ax - 15, y + 6), (ax + 15, y + 6), (ax + 11, y + 14), (ax - 11, y + 14)], fill=orange, outline=black)
        draw.line((ax - void_half, concrete_top, ax - void_half, void_bottom), fill=cyan, width=2)
        draw.line((ax + void_half, concrete_top, ax + void_half, void_bottom), fill=cyan, width=2)
        draw.line((ax - void_half, void_bottom, ax + void_half, void_bottom), fill=cyan, width=2)
        draw.rectangle((ax - void_half, void_top, ax + void_half, void_bottom), fill="#FFFFFF", outline=cyan, width=2)
        _arrow(draw, (ax, y + 2), (ax, y - 30), blue, 12)
        draw.text((ax + 18, y - 26), "Nua", fill=blue, font=small)
        return

    concrete_top, concrete_bottom, plate_top, _plate_bottom = _draw_report_single_base(draw, x, y, bw, bh)
    ax = x + bw * 0.5

    if mode == "tension_breakout":
        cone_left = x + bw * 0.18
        cone_right = x + bw * 0.82
        cone_tip = (ax, concrete_bottom - 4)
        draw.polygon((cone_left, concrete_top, cone_tip[0], cone_tip[1], cone_right, concrete_top), fill=grey, outline="#64748B")
        draw.line((cone_left, concrete_top, cone_tip[0], cone_tip[1]), fill=cyan, width=2)
        draw.line((cone_right, concrete_top, cone_tip[0], cone_tip[1]), fill=cyan, width=2)
        _draw_report_bolt_with_nut(draw, ax, plate_top, cone_tip[1], rod_color=orange)
        _arrow(draw, (ax, y + 2), (ax, y - 30), blue, 12)
        draw.text((ax + 18, y - 26), "Nua", fill=blue, font=small)
        return

    if mode == "pryout":
        anchor_bottom = (ax, concrete_bottom - 14)
        left_top = x + bw * 0.12
        small_right_top = x + bw * 0.72
        draw.polygon((left_top, concrete_top, anchor_bottom[0], anchor_bottom[1], ax, concrete_top), fill=grey, outline="#64748B")
        draw.polygon((ax, y + 72, small_right_top, concrete_top, ax, concrete_top), fill="#B8BFC7", outline="#64748B")
        draw.line((left_top, concrete_top, anchor_bottom[0], anchor_bottom[1]), fill="#FFFFFF", width=3)
        draw.line((ax, y + 72, small_right_top, concrete_top), fill="#FFFFFF", width=2)
        _draw_report_bolt_with_nut(draw, ax, plate_top, anchor_bottom[1], rod_color=orange)
        _arrow(draw, (ax + 18, plate_top - 8), (ax + 84, plate_top - 8), green, 12)
        draw.text((ax + 88, plate_top - 24), "Vua", fill=green, font=small)
        return

    _draw_report_bolt_with_nut(draw, ax, plate_top, concrete_bottom - 14, rod_color=orange)
    if mode == "steel_tension":
        _arrow(draw, (ax, y + 0), (ax, y - 32), blue, 12)
        draw.text((ax + 18, y - 28), "Nua", fill=blue, font=small)
        draw.line((ax - 18, y + 70, ax + 18, y + 52), fill=red, width=4)
        draw.line((ax - 15, y + 50, ax + 15, y + 70), fill=red, width=4)
    elif mode == "steel_shear":
        _arrow(draw, (ax + 16, plate_top - 10), (ax + 84, plate_top - 10), green, 12)
        draw.text((ax + 88, plate_top - 26), "Vua", fill=green, font=small)
        draw.line((ax - 18, concrete_top + 14, ax + 18, concrete_top + 40), fill=red, width=4)
    elif mode == "shear_breakout":
        edge_x = x + bw
        crack_start = (ax + 4, concrete_top)
        crack_end = (edge_x, concrete_bottom - 8)
        draw.polygon((crack_start[0], crack_start[1], edge_x, concrete_top, edge_x, concrete_bottom, crack_end[0], crack_end[1]), fill=grey, outline="#64748B")
        draw.line((crack_start[0], crack_start[1], crack_end[0], crack_end[1]), fill="#FFFFFF", width=3)
        _draw_report_bolt_with_nut(draw, ax, plate_top, concrete_bottom - 20, rod_color=orange)
        _arrow(draw, (ax + 16, plate_top - 10), (ax + 84, plate_top - 10), green, 12)
        draw.text((ax + 88, plate_top - 26), "Vua", fill=green, font=small)


def _draw_report_local_bearing(draw: ImageDraw.ImageDraw, x: float, y: float, bw: float, bh: float, results: AnchorResults, small) -> None:
    black = "#111827"
    blue = "#2563EB"
    orange = "#F97316"
    concrete_top, concrete_bottom, plate_top, plate_bottom = _draw_report_single_base(draw, x, y, bw, bh)
    draw.polygon((x + bw * 0.28, plate_bottom, x + bw * 0.72, plate_bottom, x + bw * 0.72, y + 100), fill="#9CA3AF", outline="#64748B")
    for idx in range(6):
        px = x + bw * 0.34 + idx * (bw * 0.32 / 5)
        py2 = plate_bottom + 14 + idx * 7
        _arrow(draw, (px, py2 + 18), (px, py2), "#FFFFFF", 7)
    for bx in (x + bw * 0.36, x + bw * 0.64):
        _draw_report_bolt_with_nut(draw, bx, plate_top, concrete_bottom - 18, rod_color=orange)
    _arrow(draw, (x + bw / 2, y + 4), (x + bw / 2, plate_bottom - 4), blue, 12)
    draw.text((x + bw + 24, y + 42), "sigma_max <= fc\nConcrete local bearing", fill=black, font=small)


def _draw_report_friction(draw: ImageDraw.ImageDraw, x: float, y: float, bw: float, bh: float, results: AnchorResults, small) -> None:
    black = "#111827"
    blue = "#2563EB"
    green = "#16A34A"
    orange = "#F97316"
    concrete_top, concrete_bottom, plate_top, _plate_bottom = _draw_report_single_base(draw, x, y, bw, bh)
    for bx in (x + bw / 2 - 58, x + bw / 2 + 58):
        _draw_report_bolt_with_nut(draw, bx, plate_top, concrete_bottom - 18, rod_color=orange)
    _arrow(draw, (x + bw / 2 - 54, y + 4), (x + bw / 2 + 46, y + 4), blue, 12)
    draw.text((x + bw / 2 - 70, y - 14), "V", fill=blue, font=small)
    _arrow(draw, (x + bw / 2 + 52, concrete_top + 54), (x + bw / 2 - 50, concrete_top + 54), green, 12)
    draw.text((x + bw / 2 + 58, concrete_top + 42), "Vfb", fill=green, font=small)
    vfb = results.values["friction_force"]
    ratio = results.inputs.v / vfb if vfb > 0 else 0.0
    draw.text((x + bw + 28, y + 42), f"V = Vfb + 2Vua\nV/Vfb = {ratio:.3f}\nVua = {results.values['vua']:.3f} kN", fill=black, font=small)


def _draw_report_tension_breakout(draw: ImageDraw.ImageDraw, x: float, y: float, bw: float, bh: float, results: AnchorResults, small) -> None:
    _draw_report_anchor_section(draw, x, y + 12, 220, 118, "tension_breakout", results, small)
    x2 = x + 260
    bw2 = 250
    black = "#111827"
    blue = "#2563EB"
    orange = "#F97316"
    red = "#DC2626"
    cyan = "#06B6D4"
    concrete_top, concrete_bottom, plate_top, _plate_bottom = _draw_report_single_base(draw, x2, y + 12, bw2, 128)
    left_anchor = x2 + bw2 * 0.38
    right_anchor = x2 + bw2 * 0.62
    bottom_y = y + 118
    cone_left = x2 + bw2 * 0.08
    cone_right = x2 + bw2 * 0.92
    draw.polygon((cone_left, concrete_top, left_anchor, bottom_y, right_anchor, bottom_y, cone_right, concrete_top), fill="#9CA3AF", outline="#64748B")
    draw.line((cone_left, concrete_top, left_anchor, bottom_y), fill=cyan, width=2)
    draw.line((left_anchor, bottom_y, right_anchor, bottom_y), fill=cyan, width=2)
    draw.line((right_anchor, bottom_y, cone_right, concrete_top), fill=cyan, width=2)
    mesh_y = concrete_top + 8
    draw.line((x2 + 8, mesh_y, x2 + bw2 - 8, mesh_y), fill=blue, width=4)
    dot_positions = (x2 + 18, x2 + 58, left_anchor + 12, right_anchor - 12, x2 + bw2 - 58, x2 + bw2 - 18)
    for dot_x in dot_positions:
        draw.ellipse((dot_x - 5, mesh_y + 6, dot_x + 5, mesh_y + 16), fill=blue, outline=blue)
    for ax in (left_anchor, right_anchor):
        _draw_report_bolt_with_nut(draw, ax, plate_top, bottom_y - 4, rod_color=orange)
    _arrow(draw, ((left_anchor + right_anchor) / 2, y + 20), ((left_anchor + right_anchor) / 2, y - 8), blue, 12)
    draw.text(((left_anchor + right_anchor) / 2 + 18, y - 6), "Nua", fill=blue, font=small)
    u_left = x2 + bw2 * 0.28
    u_right = x2 + bw2 * 0.72
    u_top = mesh_y
    u_bottom = concrete_bottom - 6
    radius = 18
    draw.line((u_left, u_bottom, u_left, u_top + radius), fill=red, width=5)
    draw.arc((u_left, u_top, u_left + 2 * radius, u_top + 2 * radius), start=180, end=270, fill=red, width=5)
    draw.line((u_left + radius, u_top, u_right - radius, u_top), fill=red, width=5)
    draw.arc((u_right - 2 * radius, u_top, u_right, u_top + 2 * radius), start=270, end=360, fill=red, width=5)
    draw.line((u_right, u_top + radius, u_right, u_bottom), fill=red, width=5)
    area = results.tension_rebar_area
    draw.text((x2 + 12, y + 150), f"Anchor reinforcement: As,N,req = {area:.2f} mm2", fill=black, font=_font(14))


def _draw_report_shear_breakout(draw: ImageDraw.ImageDraw, x: float, y: float, bw: float, bh: float, results: AnchorResults, small) -> None:
    _draw_report_anchor_section(draw, x, y + 10, 220, 122, "shear_breakout", results, small)
    x2 = x + 260
    bw2 = 250
    black = "#111827"
    green = "#16A34A"
    orange = "#F97316"
    red = "#DC2626"
    cyan = "#06B6D4"
    concrete_top, concrete_bottom, plate_top, _plate_bottom = _draw_report_single_base(draw, x2, y + 12, bw2, 128)
    ax = x2 + bw2 * 0.42
    edge_x = x2 + bw2
    crack_start = (ax + 4, concrete_top)
    crack_end = (edge_x, concrete_bottom - 8)
    draw.polygon((crack_start[0], crack_start[1], edge_x, concrete_top, edge_x, concrete_bottom, crack_end[0], crack_end[1]), fill="#9CA3AF", outline="#64748B")
    draw.line((crack_start[0], crack_start[1], crack_end[0], crack_end[1]), fill=cyan, width=2)
    _draw_report_bolt_with_nut(draw, ax, plate_top, concrete_bottom - 28, rod_color=orange)
    _arrow(draw, (ax + 26, plate_top - 8), (ax + 86, plate_top - 8), green, 12)
    draw.text((ax + 90, plate_top - 24), "Vua", fill=green, font=small)
    rebar_y = concrete_top + 34
    rebar_right = x2 + bw2 - 22
    radius = 16
    draw.line((x2 + 18, rebar_y, rebar_right - radius, rebar_y), fill=red, width=5)
    draw.arc((rebar_right - 2 * radius, rebar_y, rebar_right, rebar_y + 2 * radius), start=270, end=360, fill=red, width=5)
    draw.line((rebar_right, rebar_y + radius, rebar_right, concrete_bottom - 16), fill=red, width=5)
    dot_x = rebar_right - 10
    dot_y = rebar_y + radius - 2
    draw.ellipse((dot_x - 5, dot_y - 5, dot_x + 5, dot_y + 5), fill="#22C55E", outline="#22C55E")
    area = results.shear_rebar_area
    draw.text((x2 + 12, y + 150), f"Anchor reinforcement: As,V,req = {area:.2f} mm2", fill=black, font=_font(14))


def _draw_interaction_plot(
    draw: ImageDraw.ImageDraw,
    x: float,
    y: float,
    w: float,
    h: float,
    results: AnchorResults,
    font,
    small,
    show_point: bool = True,
) -> None:
    black = "#111827"
    blue = "#2563EB"
    grey = "#475569"
    draw.rounded_rectangle((x, y, x + w, y + h), radius=10, fill="#F8FAFC", outline="#CBD5E1", width=2)
    draw.text((x + 18, y + 16), "ACI 318-19 17.8 interaction check", fill=black, font=font)
    lines = [
        "17.8.2: If eta_N <= 0.2 or eta_V <= 0.2, interaction may be neglected.",
        "17.8.3: Otherwise, eta_N + eta_V <= 1.2.",
        f"Current eta_N = {results.governing_tension_ratio:.3f}",
        f"Current eta_V = {results.governing_shear_ratio:.3f}",
    ]
    if results.values.get("interaction_sum") is not None:
        lines.append(f"Current sum = {results.values['interaction_sum']:.3f} <= 1.2")
    elif results.values["vua"] <= 1.0e-9:
        lines.append("Current result: not applicable; anchor shear demand is zero.")
    elif not results.values.get("interaction_same_anchor_group", True):
        lines.append("Current result: not applicable; tension and shear act on different rows.")
    else:
        lines.append("Current result: interaction may be neglected.")
    for idx, line in enumerate(lines):
        fill = blue if idx >= 2 else grey
        draw.text((x + 28, y + 58 + idx * 34), line, fill=fill, font=small)
    if show_point:
        draw.ellipse((x + w - 36, y + 22, x + w - 18, y + 40), fill=blue, outline=black, width=2)


def _make_report_diagram(results: AnchorResults) -> Path:
    path = Path(tempfile.gettempdir()) / "anchor_bolt_report_diagram.png"
    width, height = 1400, 720
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    font = _font(28)
    small = _font(22)
    bold = _font(30)

    concrete = "#D1D5DB"
    concrete_dark = "#9CA3AF"
    plate = "#CBD5E1"
    orange = "#F97316"
    red = "#EF4444"
    blue = "#2563EB"
    green = "#16A34A"
    black = "#111827"
    white = "#FFFFFF"

    i = results.inputs

    # Plan view
    origin_x, origin_y = 90, 105
    plan_scale = min(520 / i.pedestal_b, 520 / i.pedestal_l)
    ped_w, ped_h = i.pedestal_b * plan_scale, i.pedestal_l * plan_scale
    plate_w, plate_h = i.plate_b * plan_scale, i.plate_l * plan_scale
    cx, cy = origin_x + ped_w / 2, origin_y + ped_h / 2
    draw.rectangle((origin_x, origin_y, origin_x + ped_w, origin_y + ped_h), fill=concrete, outline=black, width=3)
    draw.rectangle((cx - plate_w / 2, cy - plate_h / 2, cx + plate_w / 2, cy + plate_h / 2), fill=plate, outline=black, width=3)
    ax1, ax2 = cx - i.s2 * plan_scale / 2, cx + i.s2 * plan_scale / 2
    ay1, ay2 = cy - i.s1 * plan_scale / 2, cy + i.s1 * plan_scale / 2
    draw.line((cx, origin_y + 20, cx, origin_y + ped_h - 20), fill=green, width=3)
    draw.line((origin_x + 20, cy, origin_x + ped_w - 20, cy), fill=green, width=3)
    selected_y = ay2 if i.shear_case == 1 else ay1
    released_y = None if i.shear_case == 1 else ay2
    selected_label = "Case 1: near row bears in round holes" if i.shear_case == 1 else "Case 2: far row bears; near row slotted"
    edge_y = origin_y + ped_h
    edge_distance = max(edge_y - selected_y, 10)
    spread = 1.5 * edge_distance
    x_left = max(origin_x, ax1 - spread)
    x_right = min(origin_x + ped_w, ax2 + spread)
    draw.line((ax1, selected_y, x_left, edge_y), fill=white, width=5)
    draw.line((ax2, selected_y, x_right, edge_y), fill=white, width=5)
    draw.line((x_left, edge_y, x_right, edge_y), fill=white, width=5)
    draw.line((origin_x + 15, selected_y, origin_x + ped_w - 15, selected_y), fill=blue, width=3)
    for y in (ay1, ay2):
        for x in (ax1, ax2):
            is_released = released_y is not None and abs(y - released_y) < 0.1
            is_resisting = abs(y - selected_y) < 0.1
            _draw_plan_anchor_hole(draw, x, y, slotted=is_released, resisting=is_resisting, bolt_color=red)
    _arrow(draw, (cx, selected_y - 80), (cx, selected_y + 80), blue, 14)
    draw.text((cx + 24, selected_y + 16), "V", fill=blue, font=font)
    moment_r = max(42, min(70, plate_w * 0.16, plate_h * 0.16))
    _curved_moment(draw, (cx - moment_r * 0.6, cy + moment_r * 0.15), int(moment_r), blue, font)
    draw.text((origin_x, origin_y + ped_h + 126), selected_label, fill=blue, font=small)
    if released_y is not None:
        draw.text((cx - 130, released_y + 36), "slotted holes release shear", fill="#64748B", font=_font(18))
    draw.text((origin_x, 48), "Plan View", fill=black, font=bold)
    plate_left = cx - plate_w / 2
    plate_right = cx + plate_w / 2
    plate_top = cy - plate_h / 2
    plate_bottom = cy + plate_h / 2
    _dimension(draw, (origin_x, origin_y + ped_h + 52), (origin_x + ped_w, origin_y + ped_h + 52), _dim_label("B", i.pedestal_b), small)
    _dimension(draw, (origin_x + ped_w + 45, origin_y), (origin_x + ped_w + 45, origin_y + ped_h), _dim_label("L", i.pedestal_l), small, vertical=True, image=img)
    _dimension(draw, (ax1, origin_y + ped_h + 22), (ax2, origin_y + ped_h + 22), _dim_label("s2", i.s2), small)
    _dimension(draw, (origin_x + ped_w + 150, ay1), (origin_x + ped_w + 150, ay2), _dim_label("s1", i.s1), small, vertical=True, image=img)
    _dimension(draw, (plate_left, plate_top - 34), (plate_right, plate_top - 34), _dim_label("b", i.plate_b), small)
    _dimension(draw, (plate_left - 42, plate_top), (plate_left - 42, plate_bottom), _dim_label("l", i.plate_l), small, vertical=True, image=img)
    draw.text((origin_x + 18, origin_y + 18), "Concrete pedestal", fill=black, font=small)
    draw.text((plate_left + 18, plate_top + 18), "Steel plate", fill=black, font=small)

    # Elevation view
    ex, ey = 780, 105
    elev_scale = min(370 / i.pedestal_b, 440 / i.ha)
    pedestal_w, pedestal_h = i.pedestal_b * elev_scale, i.ha * elev_scale
    surface_y = ey + 75
    bottom_y = surface_y + pedestal_h
    steel_w = i.plate_b * elev_scale
    steel_h = 16
    steel_x1 = ex + pedestal_w / 2 - steel_w / 2
    steel_x2 = ex + pedestal_w / 2 + steel_w / 2
    draw.rectangle((ex, surface_y, ex + pedestal_w, bottom_y), fill=concrete, outline=black, width=3)
    draw.rectangle((steel_x1, surface_y - steel_h, steel_x2, surface_y), fill=concrete_dark, outline=black, width=2)
    elev_s2 = i.s2 * elev_scale
    embed_len = min(i.hef * elev_scale, pedestal_h)
    hook_len = max(24.0, min(54.0, i.eh * elev_scale))
    hook_radius = max(12.0, min(20.0, hook_len * 0.35))
    for idx, x in enumerate((ex + pedestal_w / 2 - elev_s2 / 2, ex + pedestal_w / 2 + elev_s2 / 2)):
        hook_dir = -1 if idx == 0 else 1
        _draw_elevation_anchor(
            draw,
            x=x,
            plate_top_y=surface_y - steel_h,
            hook_y=surface_y + embed_len,
            hook_len=hook_len,
            hook_radius=hook_radius,
            hook_dir=hook_dir,
            shaft_width=14,
            nut_color=orange,
            outline=black,
            washer_fill="#CBD5E1",
        )
    draw.line((ex + pedestal_w / 2, ey + 30, ex + pedestal_w / 2, bottom_y + 20), fill=green, width=3)
    _arrow(draw, (ex + pedestal_w / 2, ey + 30), (ex + pedestal_w / 2, surface_y + 30), blue, 14)
    draw.text((ex + pedestal_w / 2 + 25, ey + 40), "N", fill=blue, font=font)
    _dimension(draw, (ex + pedestal_w + 70, surface_y), (ex + pedestal_w + 70, surface_y + embed_len), _dim_label("hef", i.hef), small, vertical=True, image=img)
    _dimension(draw, (ex + pedestal_w + 120, surface_y), (ex + pedestal_w + 120, bottom_y), _dim_label("ha", i.ha), small, vertical=True, image=img)
    _dimension(draw, (steel_x1, surface_y - steel_h - 58), (steel_x2, surface_y - steel_h - 58), _dim_label("b", i.plate_b), small)
    draw.text((ex, 48), "Elevation View", fill=black, font=bold)
    draw.text((ex + 28, surface_y + pedestal_h * 0.52), "Concrete pedestal", fill=black, font=small)
    draw.text((steel_x1, surface_y - steel_h - 88), "Steel plate", fill=black, font=small)

    # Result ribbon
    ribbon_y = 650
    status = results.overall_status
    status_color = {"OK": "#16A34A", "Rebar Required": "#D97706", "NG": "#DC2626"}.get(status, "#475569")
    draw.rounded_rectangle((760, ribbon_y, 1310, ribbon_y + 48), radius=12, fill="#F8FAFC", outline="#CBD5E1", width=2)
    draw.text((780, ribbon_y + 10), f"Overall status: {status}", fill=status_color, font=bold)

    img.save(path)
    return path


def _dimension(draw: ImageDraw.ImageDraw, p1, p2, label: str, font, vertical: bool = False, image: Image.Image | None = None) -> None:
    black = "#111827"
    draw.line((p1, p2), fill=black, width=2)
    if vertical:
        x, y1 = p1
        _, y2 = p2
        draw.line((x - 10, y1, x + 10, y1), fill=black, width=2)
        draw.line((x - 10, y2, x + 10, y2), fill=black, width=2)
        draw.polygon([(x, y1), (x - 6, y1 + 16), (x + 6, y1 + 16)], fill=black)
        draw.polygon([(x, y2), (x - 6, y2 - 16), (x + 6, y2 - 16)], fill=black)
        if image is not None:
            _draw_rotated_text(image, (x + 20, (y1 + y2) / 2), label, font, black)
        else:
            draw.text((x + 12, (y1 + y2) / 2 - 13), label, fill=black, font=font)
    else:
        x1, y = p1
        x2, _ = p2
        draw.line((x1, y - 10, x1, y + 10), fill=black, width=2)
        draw.line((x2, y - 10, x2, y + 10), fill=black, width=2)
        draw.polygon([(x1, y), (x1 + 16, y - 6), (x1 + 16, y + 6)], fill=black)
        draw.polygon([(x2, y), (x2 - 16, y - 6), (x2 - 16, y + 6)], fill=black)
        draw.text(((x1 + x2) / 2 - 16, y + 12), label, fill=black, font=font)


def _draw_plan_anchor_hole(
    draw: ImageDraw.ImageDraw,
    x: float,
    y: float,
    *,
    slotted: bool,
    resisting: bool,
    bolt_color: str,
) -> None:
    black = "#111827"
    blue = "#2563EB"
    if slotted:
        slot_w = 28
        slot_len = 58
        draw.rounded_rectangle(
            (x - slot_w / 2, y - slot_len / 2, x + slot_w / 2, y + slot_len / 2),
            radius=slot_w / 2,
            fill="#FFFFFF",
            outline=black,
            width=3,
        )
    else:
        draw.ellipse((x - 17, y - 17, x + 17, y + 17), fill="#FFFFFF", outline=black, width=3)
    if resisting:
        draw.ellipse((x - 25, y - 25, x + 25, y + 25), outline=blue, width=5)
    draw.ellipse((x - 9, y - 9, x + 9, y + 9), fill=bolt_color, outline=black, width=2)


def _draw_elevation_anchor(
    draw: ImageDraw.ImageDraw,
    *,
    x: float,
    plate_top_y: float,
    hook_y: float,
    hook_len: float,
    hook_radius: float,
    hook_dir: int,
    shaft_width: int,
    nut_color: str,
    outline: str,
    washer_fill: str,
) -> None:
    shaft_top = plate_top_y - 34
    draw.line((x, shaft_top, x, hook_y - hook_radius), fill=outline, width=shaft_width)
    if hook_dir < 0:
        center_x = x - hook_radius
        angle_start, angle_end = 0.0, math.pi / 2.0
    else:
        center_x = x + hook_radius
        angle_start, angle_end = math.pi, math.pi / 2.0
    center_y = hook_y - hook_radius
    arc_points = []
    for step in range(13):
        t = step / 12
        angle = angle_start + (angle_end - angle_start) * t
        arc_points.append((center_x + hook_radius * math.cos(angle), center_y + hook_radius * math.sin(angle)))
    draw.line(arc_points, fill=outline, width=shaft_width, joint="curve")
    draw.line((x + hook_dir * hook_radius, hook_y, x + hook_dir * hook_len, hook_y), fill=outline, width=shaft_width)

    washer_w, washer_h = 28, 7
    washer_y1 = plate_top_y - washer_h
    draw.rectangle((x - washer_w / 2, washer_y1, x + washer_w / 2, plate_top_y), fill=washer_fill, outline=outline, width=2)

    nut_w, nut_h = 24, 18
    nut_top = washer_y1 - nut_h
    draw.rectangle((x - 5, nut_top - 16, x + 5, nut_top), fill=nut_color, outline=outline)
    draw.rectangle((x - nut_w * 0.32, nut_top, x + nut_w * 0.32, washer_y1), fill=nut_color, outline=outline)
    draw.polygon(
        [
            (x - nut_w / 2, washer_y1 - 2),
            (x + nut_w / 2, washer_y1 - 2),
            (x + nut_w * 0.38, washer_y1 + 6),
            (x - nut_w * 0.38, washer_y1 + 6),
        ],
        fill=nut_color,
        outline=outline,
    )


def _draw_rotated_text(image: Image.Image, center, text: str, font, fill: str) -> None:
    bbox = ImageDraw.Draw(image).textbbox((0, 0), text, font=font)
    text_w = bbox[2] - bbox[0]
    text_h = bbox[3] - bbox[1]
    label_img = Image.new("RGBA", (text_w + 10, text_h + 10), (255, 255, 255, 0))
    label_draw = ImageDraw.Draw(label_img)
    label_draw.text((5 - bbox[0], 5 - bbox[1]), text, fill=fill, font=font)
    rotated = label_img.rotate(90, expand=True)
    x, y = center
    image.paste(rotated, (int(x - rotated.width / 2), int(y - rotated.height / 2)), rotated)


def _arrow(draw: ImageDraw.ImageDraw, p1, p2, color: str, head: int) -> None:
    draw.line((p1, p2), fill=color, width=8)
    x1, y1 = p1
    x2, y2 = p2
    angle = __import__("math").atan2(y2 - y1, x2 - x1)
    left = (x2 - head * __import__("math").cos(angle - 0.55), y2 - head * __import__("math").sin(angle - 0.55))
    right = (x2 - head * __import__("math").cos(angle + 0.55), y2 - head * __import__("math").sin(angle + 0.55))
    draw.polygon([p2, left, right], fill=color)


def _curved_moment(draw: ImageDraw.ImageDraw, center, radius: int, color: str, font) -> None:
    x, y = center
    box = (x - radius, y - radius, x + radius, y + radius)
    draw.arc(box, start=90, end=270, fill=color, width=7)
    tip = (x, y + radius)
    head = 18
    draw.polygon(
        [
            tip,
            (tip[0] - head, tip[1] - head * 0.48),
            (tip[0] - head, tip[1] + head * 0.48),
        ],
        fill=color,
    )
    draw.text((x + radius * 0.45, y + radius * 0.62), "M", fill=color, font=font)


def _font(size: int):
    candidates = [
        "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/Arial.ttf",
        "C:/Windows/Fonts/msyh.ttc",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size=size)
        except Exception:
            continue
    return ImageFont.load_default()


def _dim_label(symbol: str, value: float) -> str:
    return f"{symbol}={_fmt(value, 1)}"


def _fmt(value: float, decimals: int = 3) -> str:
    if abs(value) >= 1000 and decimals > 1:
        return f"{value:,.{decimals}f}"
    return f"{value:.{decimals}f}".rstrip("0").rstrip(".")
